from __future__ import annotations
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph


@dataclass
class ParsedDocument:
    title: str
    sections: list[Section] = field(default_factory=list)
    raw_text: str = ""


@dataclass
class Section:
    heading: str
    level: int  # 1 = H1, 2 = H2, etc.
    paragraphs: list[str] = field(default_factory=list)
    # List items grouped: each sub-list is one bulleted/numbered list block
    lists: list[list[ListItem]] = field(default_factory=list)


@dataclass
class ListItem:
    text: str
    is_bold: bool = False  # bold = correct answer hint


def parse_docx(path: str | Path) -> ParsedDocument:
    doc = Document(str(path))
    title = ""
    sections: list[Section] = []
    current_section: Section | None = None
    current_list: list[ListItem] | None = None
    raw_lines: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            # Gap between list blocks — close current list
            if current_list:
                if current_section:
                    current_section.lists.append(current_list)
                current_list = None
            continue

        raw_lines.append(text)
        style_name = _style_name(para)

        # Headings
        if style_name.startswith("heading"):
            if current_list and current_section:
                current_section.lists.append(current_list)
                current_list = None

            level = _heading_level(style_name)
            if level == 1 and not title:
                title = text
            current_section = Section(heading=text, level=level)
            sections.append(current_section)

        # List paragraphs
        elif style_name in ("list paragraph", "list bullet", "list number") or _is_list(para):
            if current_section is None:
                current_section = Section(heading="", level=0)
                sections.append(current_section)
            if current_list is None:
                current_list = []
            bold = _has_bold(para)
            current_list.append(ListItem(text=text, is_bold=bold))

        # Regular paragraph
        else:
            if current_list and current_section:
                current_section.lists.append(current_list)
                current_list = None
            if current_section is None:
                current_section = Section(heading="", level=0)
                sections.append(current_section)
            current_section.paragraphs.append(text)

    # Flush any trailing list
    if current_list and current_section:
        current_section.lists.append(current_list)

    if not title and sections:
        title = sections[0].heading or "Untitled Activity"

    return ParsedDocument(
        title=title,
        sections=sections,
        raw_text="\n".join(raw_lines),
    )


def _heading_level(style_name: str) -> int:
    for i in range(1, 7):
        if str(i) in style_name:
            return i
    return 1


def _is_list(para: Paragraph) -> bool:
    """Detect list paragraphs by numPr XML element."""
    return para._element.find(
        ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr"
    ) is not None


def _has_bold(para: Paragraph) -> bool:
    return any(run.bold for run in para.runs if run.text.strip())


def _style_name(para: Paragraph) -> str:
    style = getattr(para, "style", None)
    name = getattr(style, "name", "") or ""
    return name.lower()
